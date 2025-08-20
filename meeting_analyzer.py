"""
Meeting Notes AI Analyzer - Local Version
Analyzes meeting notes and generates summaries with action items
"""

import os
import tempfile
import re
from datetime import datetime
from pathlib import Path
import argparse

# Try to import required libraries
try:
    import docx
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not available. Install with: pip install python-docx")

try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False
    print("⚠️  Ollama not available. Install with: pip install ollama")

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("⚠️  OpenAI not available. Install with: pip install openai")

class MeetingAnalyzer:
    def __init__(self, model_type="local"):
        self.model_type = model_type
        
        if model_type == "ollama" and OLLAMA_AVAILABLE:
            self.client = ollama
            print("✓ Using Ollama for local AI processing")
        elif model_type == "openai" and OPENAI_AVAILABLE:
            self.client = openai.OpenAI()
            print("✓ Using OpenAI for analysis")
        else:
            self.model_type = "local"
            print("✓ Using local text processing (no AI)")
    
    def analyze_meeting_notes(self, text):
        """Main analysis function"""
        if self.model_type == "ollama":
            return self._analyze_with_ollama(text)
        elif self.model_type == "openai":
            return self._analyze_with_openai(text)
        else:
            return self._analyze_with_local_processing(text)
    
    def _analyze_with_ollama(self, text):
        """Use Ollama for AI analysis"""
        prompt = self._create_analysis_prompt(text)
        
        try:
            response = self.client.chat(
                model='llama3.1',
                messages=[{'role': 'user', 'content': prompt}]
            )
            return response['message']['content']
        except Exception as e:
            print(f"Ollama error: {e}")
            return self._analyze_with_local_processing(text)
    
    def _analyze_with_openai(self, text):
        """Use OpenAI for analysis"""
        prompt = self._create_analysis_prompt(text)
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert meeting analyst."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"OpenAI error: {e}")
            return self._analyze_with_local_processing(text)
    
    def _analyze_with_local_processing(self, text):
        """Fallback local text processing without AI"""
        print("Using local text processing...")
        
        # Extract basic information
        lines = text.split('\n')
        word_count = len(text.split())
        
        # Simple keyword extraction for action items
        action_keywords = ['action', 'todo', 'task', 'assign', 'responsible', 'deadline', 'due']
        action_lines = []
        
        for line in lines:
            if any(keyword in line.lower() for keyword in action_keywords):
                action_lines.append(line.strip())
        
        # Extract names (simple capitalized words)
        names = set()
        for word in text.split():
            if word.istitle() and len(word) > 2:
                names.add(word)
        
        # Extract dates (simple pattern)
        date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b'
        dates = re.findall(date_pattern, text)
        
        # Generate summary
        summary = f"""
# Meeting Summary Report

**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M')}
**Word Count:** {word_count} words

## Executive Summary
This meeting covered {len(lines)} discussion points with {len(action_lines)} potential action items identified.

## Key Participants
{', '.join(list(names)[:10]) if names else 'Names not clearly identified'}

## Potential Action Items
"""
        
        for i, action in enumerate(action_lines[:10], 1):
            summary += f"{i}. {action}\n"
        
        if not action_lines:
            summary += "No clear action items detected in the text.\n"
        
        summary += f"""
## Important Dates Mentioned
{', '.join(dates) if dates else 'No specific dates found'}

## Next Steps
- Review and validate the extracted action items
- Assign specific owners and deadlines
- Schedule follow-up meeting if needed

## Note
This summary was generated using local text processing. For better AI-powered analysis, install Ollama or configure OpenAI API.
"""
        
        return summary
    
    def _create_analysis_prompt(self, text):
        return f"""
        Analyze the following meeting notes and provide a comprehensive summary:
        
        **Meeting Notes:**
        {text}
        
        **Please provide:**
        
        1. **Executive Summary** (2-3 sentences)
        2. **Key Discussion Points** (bullet points)
        3. **Action Items** (format: Task | Assignee | Due Date | Priority)
        4. **Decisions Made** (clear decisions reached)
        5. **Risks/Concerns** (if any mentioned)
        6. **Next Steps** (immediate follow-ups)
        7. **Meeting Metadata** (extract date, attendees, duration if mentioned)
        
        Format the response in clean markdown for easy conversion to Word document.
        """
    
    def generate_word_document(self, analysis_text, meeting_title="Meeting Summary"):
        """Convert analysis to Word document"""
        if not DOCX_AVAILABLE:
            # Save as text file instead
            return self._save_as_text(analysis_text, meeting_title)
        
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading(meeting_title, 0)
        
        # Add date
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("=" * 50)
        
        # Process and add content
        lines = analysis_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line.startswith('**') and line.endswith('**'):
                # This is a heading
                heading_text = line.replace('**', '')
                doc.add_heading(heading_text, level=1)
            elif line.startswith('# '):
                # Markdown heading
                heading_text = line.replace('# ', '')
                doc.add_heading(heading_text, level=1)
            elif line.startswith('## '):
                # Markdown subheading
                heading_text = line.replace('## ', '')
                doc.add_heading(heading_text, level=2)
            elif line.startswith('- ') or line.startswith('* '):
                # This is a bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                # Numbered list
                doc.add_paragraph(line, style='List Number')
            elif line and not line.startswith('#'):
                # Regular paragraph
                doc.add_paragraph(line)
        
        return doc
    
    def _save_as_text(self, analysis_text, meeting_title):
        """Fallback: save as text file"""
        class TextDoc:
            def __init__(self, content, title):
                self.content = f"{title}\n{'='*len(title)}\n\n{content}"
            
            def save(self, filename):
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(self.content)
        
        return TextDoc(analysis_text, meeting_title)

def main():
    parser = argparse.ArgumentParser(description="Analyze meeting notes with AI")
    parser.add_argument("input_file", nargs='?', help="Path to meeting notes text file")
    parser.add_argument("-o", "--output", help="Output document path")
    parser.add_argument("-m", "--model", choices=["local", "ollama", "openai"], 
                       default="local", help="AI model to use")
    parser.add_argument("--interactive", action="store_true", help="Interactive mode")
    
    args = parser.parse_args()
    
    # Interactive mode
    if args.interactive or not args.input_file:
        print("\n🤖 Meeting Notes AI Analyzer")
        print("=" * 40)
        
        # Get input method
        while True:
            choice = input("\nChoose input method:\n1. Load from file\n2. Paste text directly\nEnter choice (1/2): ").strip()
            
            if choice == "1":
                file_path = input("Enter path to meeting notes file: ").strip()
                if os.path.exists(file_path):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        meeting_text = f.read()
                    break
                else:
                    print("File not found. Please try again.")
                    
            elif choice == "2":
                print("Paste your meeting notes (press Ctrl+Z then Enter on Windows, or Ctrl+D on Mac/Linux to finish):")
                import sys
                meeting_text = sys.stdin.read()
                break
            else:
                print("Invalid choice. Please enter 1 or 2.")
        
        # Get model choice
        print(f"\nAvailable models:")
        print(f"- local (always available)")
        if OLLAMA_AVAILABLE:
            print(f"- ollama (AI-powered)")
        if OPENAI_AVAILABLE:
            print(f"- openai (requires API key)")
        
        model_choice = input(f"Choose model ({args.model}): ").strip() or args.model
        
    else:
        # File mode
        if not os.path.exists(args.input_file):
            print(f"Error: File '{args.input_file}' not found")
            return
        
        with open(args.input_file, 'r', encoding='utf-8') as f:
            meeting_text = f.read()
        model_choice = args.model
    
    if not meeting_text.strip():
        print("Error: No meeting notes provided")
        return
    
    # Initialize analyzer
    analyzer = MeetingAnalyzer(model_type=model_choice)
    
    print("\n🔍 Analyzing meeting notes...")
    analysis = analyzer.analyze_meeting_notes(meeting_text)
    
    # Generate output
    if args.output:
        output_path = args.output
    else:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        if DOCX_AVAILABLE:
            output_path = f"meeting_summary_{timestamp}.docx"
        else:
            output_path = f"meeting_summary_{timestamp}.txt"
    
    print("📄 Generating document...")
    doc = analyzer.generate_word_document(analysis)
    doc.save(output_path)
    
    print(f"✅ Meeting summary saved to: {output_path}")
    
    # Show preview
    print("\n📋 Preview:")
    print("-" * 50)
    print(analysis[:500] + "..." if len(analysis) > 500 else analysis)

if __name__ == "__main__":
    main()
